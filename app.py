#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import io
import random
from datetime import datetime

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import Database

# 页面配置必须在所有 Streamlit 命令之前
st.set_page_config(
    page_title="虚词练习生成器", layout="wide", initial_sidebar_state="expanded"
)

# 虚词列表
EMPTY_WORDS = [
    "而",
    "何",
    "乎",
    "乃",
    "其",
    "且",
    "若",
    "所",
    "为",
    "焉",
    "也",
    "以",
    "因",
    "于",
    "与",
    "则",
    "者",
    "之",
]


# 初始化数据库
@st.cache_resource
def get_db():
    return Database()


db = get_db()

# 主界面
st.title("虚词练习生成器")

# 筛选条件
st.markdown("### 选择虚词和题目数量")

col1, col2 = st.columns(2)
with col1:
    filter_empty_words = st.multiselect(
        "选择虚词（可多选）", EMPTY_WORDS, default=EMPTY_WORDS, key="filter_words"
    )
with col2:
    st.markdown("**题目数量**")
    question_count_option = st.radio(
        "选择题目数量",
        options=["10", "15", "20", "30", "所有"],
        horizontal=True,
        index=4,
        key="question_count_option",
    )

# 生成按钮
if st.button("下载", type="primary", use_container_width=True):
    if not filter_empty_words:
        st.error("请至少选择一个虚词")
    else:
        # 1. 获取所有符合虚词条件的句子（从 sentence 表模糊匹配）
        all_sentences = db.get_all_sentences(empty_word_filter=None)

        # 数据去重处理：基于去除标点后的文本内容去重，并处理包含关系
        import re

        def normalize_text(text):
            return re.sub(r"[^\w\u4e00-\u9fa5]", "", text)

        # 1. 先进行标准化的精确去重
        temp_unique_map = {}
        for s in all_sentences:
            norm_text = normalize_text(s["sentence"])
            if norm_text not in temp_unique_map:
                temp_unique_map[norm_text] = s
            else:
                if len(s["sentence"]) > len(temp_unique_map[norm_text]["sentence"]):
                    temp_unique_map[norm_text] = s

        # 2. 处理包含关系（子集去重）
        sorted_sentences = sorted(
            temp_unique_map.values(),
            key=lambda x: len(normalize_text(x["sentence"])),
            reverse=True,
        )

        final_unique_sentences = []
        kept_normalized_texts = []

        for s in sorted_sentences:
            current_norm = normalize_text(s["sentence"])
            is_subset = False
            for kept_norm in kept_normalized_texts:
                if current_norm in kept_norm:
                    is_subset = True
                    break
            if not is_subset:
                final_unique_sentences.append(s)
                kept_normalized_texts.append(current_norm)

        unique_sentences = final_unique_sentences
        filtered_sentences = []
        target_words = set(filter_empty_words)

        for s in unique_sentences:
            found_words = [w for w in target_words if w in s["sentence"]]
            if found_words:
                selected_word = random.choice(found_words)
                filtered_sentences.append(
                    {
                        "id": s["id"],
                        "sentence": s["sentence"],
                        "empty_word": selected_word,
                        "nos": s["nos"],
                        "tags": s["tags"],
                    }
                )

        if len(filtered_sentences) == 0:
            st.error("没有符合条件的例句")
        else:
            random.shuffle(filtered_sentences)

            if question_count_option == "所有":
                selected_sentences = filtered_sentences
            else:
                question_count = int(question_count_option)
                selected_sentences = random.sample(
                    filtered_sentences, min(question_count, len(filtered_sentences))
                )

            if len(selected_sentences) == 0:
                st.error("没有可用的题目")
            else:
                try:
                    # 生成 Word 文档
                    paper_title = f"虚词练习 {datetime.now().strftime('%Y-%m-%d')}"
                    doc = Document()

                    title = doc.add_heading(paper_title, 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

                    for i, sentence_data in enumerate(selected_sentences, 1):
                        sentence = sentence_data.get("sentence")
                        empty_word = sentence_data.get("empty_word")

                        if not sentence:
                            continue

                        para = doc.add_paragraph(f"{i}. ", style="Normal")
                        para.add_run(sentence)
                        # para.add_run(f"   [{empty_word}]")

                        # doc.add_paragraph()  # 空行

                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_bytes = doc_io.getvalue()
                    doc_io.close()

                    st.success(f"已生成 {len(selected_sentences)} 道题目的试卷")

                    st.download_button(
                        "📥 下载试卷",
                        doc_bytes,
                        f"{paper_title}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="export_download",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"生成 Word 文档时出错: {str(e)}")
                    import traceback

                    st.code(traceback.format_exc())
